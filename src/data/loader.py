"""Document loaders for various file formats"""

import gc
from pathlib import Path
from typing import List, Optional, Callable
from dataclasses import dataclass, field

from src.logger import setup_logger

logger = setup_logger(__name__)

# Progress callback type
ProgressCallback = Callable[[str, int, int], None]


@dataclass
class Document:
    """Represents a loaded document"""
    content: str
    metadata: dict = field(default_factory=dict)
    
    @property
    def source(self) -> str:
        return self.metadata.get("source", "unknown")
    
    @property
    def page_count(self) -> int:
        return self.metadata.get("page_count", 1)


class PDFLoader:
    """Load PDF documents using PyMuPDF (fast) with pdfplumber fallback (tables)"""
    
    # Processing configuration
    BATCH_SIZE = 50  # Pages per batch
    LARGE_FILE_THRESHOLD_MB = 5  # Files above this use batched processing
    
    def __init__(self, extract_tables: bool = False, use_fast_mode: bool = True):
        """
        Initialize PDF loader.
        
        Args:
            extract_tables: Whether to extract tables (slower, uses pdfplumber)
            use_fast_mode: Use PyMuPDF for speed (default True)
        """
        self.extract_tables = extract_tables
        self.use_fast_mode = use_fast_mode
    
    def load(
        self, 
        file_path: str | Path,
        progress_callback: Optional[ProgressCallback] = None,
    ) -> Document:
        """
        Load a PDF file and extract text.
        
        Args:
            file_path: Path to PDF file
            progress_callback: Optional callback(filename, current_page, total_pages)
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"PDF not found: {file_path}")
        
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        
        # Use fast PyMuPDF for large files or when fast mode enabled
        if self.use_fast_mode:
            return self._load_with_pymupdf(file_path, progress_callback)
        else:
            return self._load_with_pdfplumber(file_path, progress_callback)
    
    def _load_with_pymupdf(
        self, 
        file_path: Path,
        progress_callback: Optional[ProgressCallback] = None,
    ) -> Document:
        """Fast PDF loading using PyMuPDF with batched processing."""
        import fitz  # PyMuPDF
        
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        pages_text = []
        
        doc = fitz.open(file_path)
        page_count = len(doc)
        
        logger.info(f"Processing: {file_path.name} ({file_size_mb:.1f} MB, {page_count} pages)")
        
        # Process in batches for memory efficiency
        for batch_start in range(0, page_count, self.BATCH_SIZE):
            batch_end = min(batch_start + self.BATCH_SIZE, page_count)
            
            for page_num in range(batch_start, batch_end):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    pages_text.append(f"[Page {page_num + 1}]\n{text}")
            
            # Progress callback
            if progress_callback:
                progress_callback(file_path.name, batch_end, page_count)
            else:
                logger.info(f"  → Pages {batch_start + 1}-{batch_end}/{page_count} processed...")
            
            # Free memory between batches for large files
            if file_size_mb > self.LARGE_FILE_THRESHOLD_MB:
                gc.collect()
        
        doc.close()
        
        content = "\n\n".join(pages_text)
        
        logger.info(f"  ✓ Completed: {file_path.name} ({len(content):,} chars extracted)")
        
        return Document(
            content=content,
            metadata={
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": "pdf",
                "page_count": page_count,
                "file_size_mb": round(file_size_mb, 2),
            }
        )
    
    def _load_with_pdfplumber(
        self,
        file_path: Path,
        progress_callback: Optional[ProgressCallback] = None,
    ) -> Document:
        """Load PDF with pdfplumber (slower but better table extraction)."""
        import pdfplumber
        
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        pages_text = []
        tables_text = []
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            pdf_metadata = pdf.metadata or {}
            
            logger.info(f"Processing: {file_path.name} ({file_size_mb:.1f} MB, {page_count} pages)")
            
            for batch_start in range(0, page_count, self.BATCH_SIZE):
                batch_end = min(batch_start + self.BATCH_SIZE, page_count)
                
                for page_num in range(batch_start, batch_end):
                    page = pdf.pages[page_num]
                    
                    # Extract text
                    text = page.extract_text() or ""
                    if text.strip():
                        pages_text.append(f"[Page {page_num + 1}]\n{text}")
                    
                    # Extract tables
                    if self.extract_tables:
                        tables = page.extract_tables()
                        for table in tables:
                            table_text = self._table_to_text(table)
                            if table_text:
                                tables_text.append(f"[Table on Page {page_num + 1}]\n{table_text}")
                
                # Progress callback
                if progress_callback:
                    progress_callback(file_path.name, batch_end, page_count)
                else:
                    logger.info(f"  → Pages {batch_start + 1}-{batch_end}/{page_count} processed...")
                
                # Free memory between batches
                if file_size_mb > self.LARGE_FILE_THRESHOLD_MB:
                    gc.collect()
        
        # Combine all content
        content_parts = pages_text + tables_text
        content = "\n\n".join(content_parts)
        
        logger.info(f"  ✓ Completed: {file_path.name} ({len(content):,} chars extracted)")
        
        return Document(
            content=content,
            metadata={
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": "pdf",
                "page_count": page_count,
                "file_size_mb": round(file_size_mb, 2),
                "title": pdf_metadata.get("Title", ""),
                "author": pdf_metadata.get("Author", ""),
            }
        )
    
    def _table_to_text(self, table: List[List]) -> str:
        """Convert table to readable text format"""
        if not table:
            return ""
        
        rows = []
        for row in table:
            # Clean cells and join
            cells = [str(cell).strip() if cell else "" for cell in row]
            if any(cells):  # Skip empty rows
                rows.append(" | ".join(cells))
        
        return "\n".join(rows)


class DocxLoader:
    """Load Word documents"""
    
    def load(self, file_path: str | Path) -> Document:
        """Load a DOCX file and extract text"""
        from docx import Document as DocxDocument
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX not found: {file_path}")
        
        doc = DocxDocument(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract tables
        tables_text = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(" | ".join(cells))
            if table_rows:
                tables_text.append("\n".join(table_rows))
        
        content = "\n\n".join(paragraphs + tables_text)
        
        return Document(
            content=content,
            metadata={
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": "docx",
                "page_count": 1,  # DOCX doesn't have explicit pages
            }
        )


class TextLoader:
    """Load plain text files"""
    
    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding
    
    def load(self, file_path: str | Path) -> Document:
        """Load a text file"""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        content = file_path.read_text(encoding=self.encoding)
        
        return Document(
            content=content,
            metadata={
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": file_path.suffix.lstrip("."),
                "page_count": 1,
            }
        )


class DocumentLoader:
    """Unified document loader - auto-detects file type"""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}
    
    def __init__(self, use_fast_pdf: bool = True, extract_tables: bool = False):
        """
        Initialize document loader.
        
        Args:
            use_fast_pdf: Use PyMuPDF for fast PDF processing (default True)
            extract_tables: Extract tables from PDFs (slower, default False)
        """
        self.pdf_loader = PDFLoader(extract_tables=extract_tables, use_fast_mode=use_fast_pdf)
        self.docx_loader = DocxLoader()
        self.text_loader = TextLoader()
    
    def load(
        self, 
        file_path: str | Path,
        progress_callback: Optional[ProgressCallback] = None,
    ) -> Document:
        """Load a document based on file extension"""
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        
        if ext == ".pdf":
            return self.pdf_loader.load(file_path, progress_callback=progress_callback)
        elif ext in {".docx", ".doc"}:
            return self.docx_loader.load(file_path)
        elif ext in {".txt", ".md"}:
            return self.text_loader.load(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def load_file(self, file_path: str | Path) -> Document:
        """Alias for load() for backward compatibility"""
        return self.load(file_path)
    
    def load_directory(
        self, 
        directory: str | Path, 
        recursive: bool = True,
        progress_callback: Optional[ProgressCallback] = None,
    ) -> List[Document]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory: Directory path
            recursive: Search subdirectories
            progress_callback: Optional callback for progress updates
        """
        directory = Path(directory)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        # Collect all files first
        pattern = "**/*" if recursive else "*"
        files = [
            f for f in directory.glob(pattern) 
            if f.is_file() and f.suffix.lower() in self.SUPPORTED_EXTENSIONS
        ]
        
        # Sort by size (smallest first for quick wins)
        files.sort(key=lambda f: f.stat().st_size)
        
        total_files = len(files)
        total_size_mb = sum(f.stat().st_size for f in files) / (1024 * 1024)
        
        logger.info(f"Found {total_files} documents ({total_size_mb:.1f} MB total)")
        logger.info("=" * 50)
        
        documents = []
        for idx, file_path in enumerate(files, 1):
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            logger.info(f"\n[{idx}/{total_files}] Loading: {file_path.name} ({file_size_mb:.1f} MB)")
            
            try:
                doc = self.load(file_path, progress_callback=progress_callback)
                documents.append(doc)
                logger.info(f"  ✓ Success: {len(doc.content):,} chars")
            except Exception as e:
                logger.error(f"  ✗ Failed: {e}")
        
        logger.info("=" * 50)
        logger.info(f"Loaded {len(documents)}/{total_files} documents from {directory}")
        return documents
